    import tkinter as tk
    from tkinter import filedialog, messagebox
    import os
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

    class Windows95Converter:
        def __init__(self, rootwordtoPDFFF):
            self.rootwordtoPDFFF = rootwordtoPDFFF
            self.rootwordtoPDFFF.title("Word to PDF Converter")
            self.rootwordtoPDFFF.geometry("500x450")
            self.rootwordtoPDFFF.resizable(False, False)
            
            # Windows 95 color scheme
            self.bg_color = "#c0c0c0"
            self.button_color = "#c0c0c0"
            self.white = "#ffffff"
            self.black = "#000000"
            self.dark_gray = "#808080"
            self.light_gray = "#dfdfdf"
            
            self.rootwordtoPDFFF.configure(bg=self.bg_color)
            self.input_file = ""
            self.output_file = ""
            
            self.create_widgets()
        
        def create_3d_frame(self, parent, relief="sunken"):
            frame = tk.Frame(parent, bg=self.bg_color, relief=relief, borderwidth=2)
            return frame
        
        def create_widgets(self):
            # Title bar simulation
            title_frame = tk.Frame(self.rootwordtoPDFFF, bg="#000080", height=25)
            title_frame.pack(fill=tk.X, padx=2, pady=2)
            title_frame.pack_propagate(False)
            
            title_label = tk.Label(title_frame, text="Word to PDF Converter", 
                                  bg="#000080", fg=self.white, 
                                  font=("MS Sans Serif", 9, "bold"))
            title_label.pack(side=tk.LEFT, padx=5)
            
            # Main container
            main_frame = self.create_3d_frame(self.rootwordtoPDFFF, relief="raised")
            main_frame.pack(fill=tk.BOTH, expand=True, padx=3, pady=3)
            
            # Header
            header_label = tk.Label(main_frame, text="Convert Word Documents to PDF", 
                                   bg=self.bg_color, fg=self.black,
                                   font=("MS Sans Serif", 10, "bold"))
            header_label.pack(pady=15)
            
            # Input file section
            input_frame = self.create_3d_frame(main_frame, relief="groove")
            input_frame.pack(fill=tk.X, padx=20, pady=10)
            
            input_label = tk.Label(input_frame, text="Input Word File:", 
                                  bg=self.bg_color, fg=self.black,
                                  font=("MS Sans Serif", 8))
            input_label.pack(anchor=tk.W, padx=10, pady=(10, 5))
            
            input_path_frame = tk.Frame(input_frame, bg=self.white, relief="sunken", borderwidth=1)
            input_path_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
            
            self.input_path_label = tk.Label(input_path_frame, text="No file selected", 
                                             bg=self.white, fg=self.black,
                                             font=("MS Sans Serif", 8), anchor=tk.W)
            self.input_path_label.pack(fill=tk.X, padx=5, pady=5)
            
            input_btn = tk.Button(input_frame, text="Browse...", 
                                 command=self.select_input_file,
                                 bg=self.button_color, fg=self.black,
                                 font=("MS Sans Serif", 8),
                                 relief="raised", borderwidth=2,
                                 width=12, height=1)
            input_btn.pack(pady=(0, 10))
            
            # Output file section
            output_frame = self.create_3d_frame(main_frame, relief="groove")
            output_frame.pack(fill=tk.X, padx=20, pady=10)
            
            output_label = tk.Label(output_frame, text="Output PDF File:", 
                                   bg=self.bg_color, fg=self.black,
                                   font=("MS Sans Serif", 8))
            output_label.pack(anchor=tk.W, padx=10, pady=(10, 5))
            
            output_path_frame = tk.Frame(output_frame, bg=self.white, relief="sunken", borderwidth=1)
            output_path_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
            
            self.output_path_label = tk.Label(output_path_frame, text="No file selected", 
                                              bg=self.white, fg=self.black,
                                              font=("MS Sans Serif", 8), anchor=tk.W)
            self.output_path_label.pack(fill=tk.X, padx=5, pady=5)
            
            output_btn = tk.Button(output_frame, text="Browse...", 
                                  command=self.select_output_file,
                                  bg=self.button_color, fg=self.black,
                                  font=("MS Sans Serif", 8),
                                  relief="raised", borderwidth=2,
                                  width=12, height=1)
            output_btn.pack(pady=(0, 10))
            
            # Convert button
            convert_btn = tk.Button(main_frame, text="Convert to PDF", 
                                   command=self.convert_file,
                                   bg=self.button_color, fg=self.black,
                                   font=("MS Sans Serif", 9, "bold"),
                                   relief="raised", borderwidth=2,
                                   width=20, height=2)
            convert_btn.pack(pady=20)
            
            # Info label
            info_label = tk.Label(main_frame, 
                                 text="Tudor Marmureanu",
                                 bg=self.bg_color, fg=self.dark_gray,
                                 font=("MS Sans Serif", 7))
            info_label.pack(pady=(0, 10))
            
            # Status bar
            status_frame = tk.Frame(self.rootwordtoPDFFF, bg=self.bg_color, relief="sunken", borderwidth=1)
            status_frame.pack(fill=tk.X, side=tk.BOTTOM, padx=2, pady=2)
            
            self.status_label = tk.Label(status_frame, text="Ready", 
                                         bg=self.bg_color, fg=self.black,
                                         font=("MS Sans Serif", 8), anchor=tk.W)
            self.status_label.pack(fill=tk.X, padx=5, pady=2)
        
        def select_input_file(self):
            filename = filedialog.askopenfilename(
                title="Select Word Document",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
            )
            if filename:
                self.input_file = filename
                self.input_path_label.config(text=os.path.basename(filename))
                self.status_label.config(text=f"Selected: {os.path.basename(filename)}")
        
        def select_output_file(self):
            filename = filedialog.asksaveasfilename(
                title="Save PDF As",
                defaultextension=".pdf",
                filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")]
            )
            if filename:
                self.output_file = filename
                self.output_path_label.config(text=os.path.basename(filename))
                self.status_label.config(text=f"Output: {os.path.basename(filename)}")
        
        def convert_docx_to_pdf(self, docx_path, pdf_path):
            """Convert DOCX to PDF using python-docx and reportlab"""
            # Read the Word document
            doc = Document(docx_path)
            
            # Create PDF
            pdf = SimpleDocTemplate(pdf_path, pagesize=letter,
                                   rightMargin=72, leftMargin=72,
                                   topMargin=72, bottomMargin=18)
            
            # Container for the 'Flowable' objects
            elements = []
            
            # Define styles
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(name='CustomBody', 
                                     parent=styles['BodyText'],
                                     fontSize=11,
                                     leading=14))
            
            # Process each paragraph
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Escape special characters for reportlab
                    text = paragraph.text.replace('&', '&amp;')
                    text = text.replace('<', '&lt;')
                    text = text.replace('>', '&gt;')
                    
                    # Determine style based on paragraph properties
                    if paragraph.style.name.startswith('Heading'):
                        style = styles['Heading1']
                    else:
                        style = styles['CustomBody']
                    
                    # Add paragraph to PDF
                    para = Paragraph(text, style)
                    elements.append(para)
                    elements.append(Spacer(1, 0.2 * inch))
            
            # Build PDF
            pdf.build(elements)
        
        def convert_file(self):
            if not self.input_file:
                messagebox.showerror("Error", "Please select an input Word file!")
                return
            
            if not self.output_file:
                messagebox.showerror("Error", "Please select an output PDF file!")
                return
            
            if not os.path.exists(self.input_file):
                messagebox.showerror("Error", "Input file does not exist!")
                return
            
            try:
                self.status_label.config(text="Converting...")
                self.rootwordtoPDFFF.update()
                
                # Convert using python-docx + reportlab
                self.convert_docx_to_pdf(self.input_file, self.output_file)
                
                self.status_label.config(text="Conversion completed successfully!")
                messagebox.showinfo("Success", 
                    "Word document converted to PDF successfully!\n\n"
                    "Note: Only text content is converted.\n"
                    "Images and complex formatting may not be preserved.")
                
            except ImportError as e:
                self.status_label.config(text="Error: Missing library!")
                missing_lib = "python-docx" if "docx" in str(e) else "reportlab"
                messagebox.showerror("Error", 
                    f"{missing_lib} library not found!\n\n"
                    "Install required libraries:\n"
                    "pip install python-docx reportlab")
            except Exception as e:
                self.status_label.config(text="Error occurred!")
                messagebox.showerror("Error", f"An error occurred:\n{str(e)}")

    def main():
        rootwordtoPDFFF = tk.Tk()
        app = Windows95Converter(rootwordtoPDFFF)
        rootwordtoPDFFF.mainloop()

    if __name__ == "__main__":
        main()
        
        
        
############################################################

    def PDFToWordConv():
        from pdf2docx import Converter
        from docx import Document
        import logging
        import io
        sys.stdout = io.StringIO()
        sys.stderr = io.StringIO()

        logging.getLogger('pdf2docx').setLevel(logging.ERROR)
        logging.basicConfig(level=logging.CRITICAL)
        logging.getLogger().setLevel(logging.CRITICAL)
        gpath = ""
        
        def openfile():
            tb.delete(0, END)
            nonlocal gpath
            filterex = [('PDF File', '*.pdf'), ('All Files', '*.*')]
            userfile = fd.askopenfile(title="Open PDF", filetypes=filterex)
            if userfile:
                tb.insert(0, userfile.name)
                gpath = userfile.name
        
        def convert_to_word():
            nonlocal gpath
            if gpath:
                try:
                    # Convert PDF to DOCX using pdf2docx
                    cv = Converter(gpath)
                    output_file = gpath.replace('.pdf', '.docx')
                    cv.convert(output_file)
                    cv.close()
                    mb.showinfo("Success", f"The file has been converted to Word.\n\nSaved as: {output_file}")
                except Exception as e:
                    mb.showerror("Error", f"Conversion failed:\n{str(e)}")
            else:
                mb.showerror("Error", "No PDF file has been selected.")
        
        def convert_to_word_and_extract_images():
            nonlocal gpath
            if gpath:
                try:
                    # Convert PDF to DOCX
                    cv = Converter(gpath)
                    output_file = "pdf.docx"
                    cv.convert(output_file)
                    cv.close()
                    
                    # Extract images from the converted DOCX
                    doc = Document(output_file)
                    folder_path = "images_from_pdf/"
                    
                    # Create folder if it doesn't exist
                    if not os.path.exists(folder_path):
                        os.makedirs(folder_path)
                    
                    imageIndex = 0
                    
                    # Extract images from document relationships
                    for rel in doc.part.rels.values():
                        if "image" in rel.target_ref:
                            try:
                                image = rel.target_part.blob
                                # Determine file extension
                                image_ext = rel.target_ref.split('.')[-1]
                                image_filename = f"{folder_path}Image_{imageIndex}.{image_ext}"
                                
                                # Save image
                                with open(image_filename, 'wb') as img_file:
                                    img_file.write(image)
                                imageIndex += 1
                            except:
                                pass
                    
                    mb.showinfo("Success", 
                               f"The file was converted to Word and {imageIndex} images were extracted.\n\n"
                               f"Word file: {output_file}\n"
                               f"Images folder: {folder_path}")
                except Exception as e:
                    mb.showerror("Error", f"Conversion failed:\n{str(e)}")
            else:
                mb.showerror("Error", "No PDF file has been selected.")
        
        W = Tk()
        W.title("PDF to Word Converter")
        W.geometry("640x400")
        W.config(bg="gray20")
        W.resizable(FALSE, FALSE)
        
        lb = Label(W, text="Select PDF file", bg="gray40", fg="cyan", bd=1)
        lb.grid(row=1, column=0, padx=10, pady=10)
        
        tb = Entry(W, width=80)
        tb.grid(row=2, column=0, padx=10)
        tb.focus()
        
        bt_browse = Button(W, text="Browse", bg="gray40", fg="cyan", width=15, bd=1, command=openfile)
        bt_browse.grid(row=2, column=1, ipady=3, padx=5)
        
        bt_convert_to_word = Button(W, text="Convert to Word", width=30, bg="gray40", fg="cyan", bd=1, command=convert_to_word)
        bt_convert_to_word.grid(row=3, column=0, columnspan=2, pady=5)
        
        bt_convert_with_images = Button(W, text="Convert to Word and extract images", width=30, bg="gray40", fg="cyan", bd=1, command=convert_to_word_and_extract_images)
        bt_convert_with_images.grid(row=4, column=0, columnspan=2, pady=5)
        
        W.mainloop()
        
    WP = tk.Button(main_frame, text="PDF to Word", command=PDFToWordConv, bg="gray40", fg="white", bd=5) #bg=black
    WP.place(x=170,y=360)