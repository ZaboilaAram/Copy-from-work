import tkinter as tk
from tkinter import filedialog, messagebox
from pdf2docx import Converter
import os
import logging
import fitz  # PyMuPDF
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logging.disable(logging.CRITICAL)

class Windows95Converter:
    def __init__(self, rootWORDDOCX):
        self.rootWORDDOCX = rootWORDDOCX
        self.rootWORDDOCX.title("PDF to Word Converter")
        self.rootWORDDOCX.geometry("500x350")
        self.rootWORDDOCX.resizable(False, False)
        
        # Windows 95 colors
        bg_color = "#c0c0c0"
        button_bg = "#c0c0c0"
        self.rootWORDDOCX.configure(bg=bg_color)
        
        # Title bar style frame
        title_frame = tk.Frame(rootWORDDOCX, bg="#000080", height=25)
        title_frame.pack(fill=tk.X)
        
        title_label = tk.Label(title_frame, text="PDF to Word Converter", 
                               bg="#000080", fg="white", font=("MS Sans Serif", 8, "bold"))
        title_label.pack(side=tk.LEFT, padx=5, pady=2)
        
        # Main content frame
        content_frame = tk.Frame(rootWORDDOCX, bg=bg_color, relief=tk.RAISED, bd=2)
        content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # File selection section
        file_frame = tk.Frame(content_frame, bg=bg_color)
        file_frame.pack(pady=20, padx=20)
        
        tk.Label(file_frame, text="Select PDF file:", bg=bg_color, 
                font=("MS Sans Serif", 8)).pack(anchor=tk.W)
        
        path_frame = tk.Frame(file_frame, bg=bg_color)
        path_frame.pack(fill=tk.X, pady=5)
        
        self.file_path = tk.StringVar()
        self.path_entry = tk.Entry(path_frame, textvariable=self.file_path, 
                                   width=40, font=("MS Sans Serif", 8))
        self.path_entry.pack(side=tk.LEFT, padx=(0, 5))
        
        browse_btn = tk.Button(path_frame, text="Browse...", 
                              command=self.browse_file,
                              bg=button_bg, relief=tk.RAISED, bd=2,
                              font=("MS Sans Serif", 8), width=10)
        browse_btn.pack(side=tk.LEFT)
        
        # Buttons section
        button_frame = tk.Frame(content_frame, bg=bg_color)
        button_frame.pack(pady=20)
        
        convert_btn = tk.Button(button_frame, text="Convert with Images (Full)", 
                               command=self.convert_file_with_images,
                               bg=button_bg, relief=tk.RAISED, bd=2,
                               font=("MS Sans Serif", 8, "bold"), 
                               width=25, height=2)
        convert_btn.pack(side=tk.LEFT, padx=5)
        
        convert_no_img_btn = tk.Button(button_frame, text="Convert without Images", 
                                       command=self.convert_file_without_images,
                                       bg=button_bg, relief=tk.RAISED, bd=2,
                                       font=("MS Sans Serif", 8, "bold"), 
                                       width=25, height=2)
        convert_no_img_btn.pack(side=tk.LEFT, padx=5)
        
        # Extract images button
        extract_frame = tk.Frame(content_frame, bg=bg_color)
        extract_frame.pack(pady=5)
        
        extract_btn = tk.Button(extract_frame, text="Extract Images to Folder", 
                               command=self.extract_images,
                               bg=button_bg, relief=tk.RAISED, bd=2,
                               font=("MS Sans Serif", 8), 
                               width=30, height=2)
        extract_btn.pack(padx=5)
        
        # Exit button
        exit_frame = tk.Frame(content_frame, bg=bg_color)
        exit_frame.pack(pady=5)
        
        exit_btn = tk.Button(exit_frame, text="Exit", 
                            command=self.rootWORDDOCX.quit,
                            bg=button_bg, relief=tk.RAISED, bd=2,
                            font=("MS Sans Serif", 8), width=15, height=1)
        exit_btn.pack(side=tk.LEFT, padx=5)
        
        # Status bar
        status_frame = tk.Frame(rootWORDDOCX, bg=bg_color, relief=tk.SUNKEN, bd=1)
        status_frame.pack(side=tk.BOTTOM, fill=tk.X)
        
        self.status_label = tk.Label(status_frame, text="Ready", 
                                     bg=bg_color, anchor=tk.W,
                                     font=("MS Sans Serif", 8))
        self.status_label.pack(side=tk.LEFT, padx=5, pady=2)
    
    def browse_file(self):
        filename = filedialog.askopenfilename(
            title="Select PDF file",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        if filename:
            self.file_path.set(filename)
            self.status_label.config(text="File selected")
    
    def convert_file_with_images(self):
        pdf_path = self.file_path.get()
        
        if not pdf_path:
            messagebox.showwarning("Warning", "Please select a PDF file first!")
            return
        
        if not os.path.exists(pdf_path):
            messagebox.showerror("Error", "The selected file does not exist!")
            return
        
        try:
            word_path = pdf_path.replace('.pdf', '.docx')
            
            self.status_label.config(text="Converting with images...")
            self.rootWORDDOCX.update()
            
            cv = Converter(pdf_path)
            cv.convert(word_path)
            cv.close()
            
            self.status_label.config(text="Conversion successful!")
            messagebox.showinfo("Success", 
                              f"File converted successfully with images!\n\nSaved as:\n{word_path}")
            
        except Exception as e:
            self.status_label.config(text="Conversion failed")
            messagebox.showerror("Error", f"Conversion failed:\n{str(e)}")
    
    def remove_images_from_docx(self, docx_path, output_path):
        """Remove all images from a DOCX file while preserving structure"""
        doc = Document(docx_path)
        
        # Remove inline images from paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Check for drawings (images)
                drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                for drawing in drawings:
                    drawing.getparent().remove(drawing)
                
                # Check for pictures
                pictures = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                for picture in pictures:
                    picture.getparent().remove(picture)
        
        # Remove images from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                            for drawing in drawings:
                                drawing.getparent().remove(drawing)
                            
                            pictures = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                            for picture in pictures:
                                picture.getparent().remove(picture)
        
        # Remove images from headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    for drawing in drawings:
                        drawing.getparent().remove(drawing)
                    
                    pictures = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                    for picture in pictures:
                        picture.getparent().remove(picture)
        
        # Remove images from footers
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                for run in paragraph.runs:
                    drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    for drawing in drawings:
                        drawing.getparent().remove(drawing)
                    
                    pictures = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                    for picture in pictures:
                        picture.getparent().remove(picture)
        
        doc.save(output_path)
    
    def convert_file_without_images(self):
        pdf_path = self.file_path.get()
        
        if not pdf_path:
            messagebox.showwarning("Warning", "Please select a PDF file first!")
            return
        
        if not os.path.exists(pdf_path):
            messagebox.showerror("Error", "The selected file does not exist!")
            return
        
        try:
            temp_path = pdf_path.replace('.pdf', '_temp.docx')
            word_path = pdf_path.replace('.pdf', '_no_images.docx')
            
            self.status_label.config(text="Converting without images...")
            self.rootWORDDOCX.update()
            
            # First, convert normally with pdf2docx to preserve structure
            cv = Converter(pdf_path)
            cv.convert(temp_path)
            cv.close()
            
            # Then remove all images from the converted document
            self.remove_images_from_docx(temp_path, word_path)
            
            # Delete temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
            
            self.status_label.config(text="Conversion successful!")
            messagebox.showinfo("Success", 
                              f"File converted successfully without images!\n\nSaved as:\n{word_path}")
            
        except Exception as e:
            self.status_label.config(text="Conversion failed")
            # Clean up temp file if exists
            temp_path = pdf_path.replace('.pdf', '_temp.docx')
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass
            messagebox.showerror("Error", f"Conversion failed:\n{str(e)}")
    
    def extract_images(self):
        pdf_path = self.file_path.get()
        
        if not pdf_path:
            messagebox.showwarning("Warning", "Please select a PDF file first!")
            return
        
        if not os.path.exists(pdf_path):
            messagebox.showerror("Error", "The selected file does not exist!")
            return
        
        try:
            # Create folder for images
            pdf_dir = os.path.dirname(pdf_path)
            pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
            images_folder = os.path.join(pdf_dir, f"{pdf_name}_images")
            
            if not os.path.exists(images_folder):
                os.makedirs(images_folder)
            
            self.status_label.config(text="Extracting images...")
            self.rootWORDDOCX.update()
            
            # Open PDF with PyMuPDF
            pdf_document = fitz.open(pdf_path)
            image_count = 0
            
            # Iterate through pages
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                images = page.get_images()
                
                # Extract each image
                for img_index, img in enumerate(images):
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # Save image
                    image_filename = f"page{page_num + 1}_img{img_index + 1}.{image_ext}"
                    image_path = os.path.join(images_folder, image_filename)
                    
                    with open(image_path, "wb") as img_file:
                        img_file.write(image_bytes)
                    
                    image_count += 1
            
            pdf_document.close()
            
            if image_count == 0:
                self.status_label.config(text="No images found!")
                messagebox.showinfo("Info", "No images found in the PDF file.")
            else:
                self.status_label.config(text=f"Extracted {image_count} images!")
                messagebox.showinfo("Success", 
                                  f"Extracted {image_count} images successfully!\n\nSaved in:\n{images_folder}")
            
        except Exception as e:
            self.status_label.config(text="Extraction failed")
            messagebox.showerror("Error", f"Image extraction failed:\n{str(e)}")

if __name__ == "__main__":
    rootWORDDOCX = tk.Tk()
    app = Windows95Converter(rootWORDDOCX)
    rootWORDDOCX.mainloop()